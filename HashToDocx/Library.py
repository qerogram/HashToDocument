# Author : qerogram

import docx, os, platform, hashlib, time
from datetime import datetime

class HashToDocument() :
    NOT_SUPPORTED_OPTION = 'Not Supported Option'

    
    def __init__(self, options = ['filename', 'size', 'md5', 'sha1', 'sha256', 'ctime', 'mtime', 'atime']) :
        #############################################
        self.path_delimiter = '/' if platform.system() == 'Darwin' else '\\'

        self.Hash_Algorithm = {
            "md5"       : hashlib.md5,
            "sha1"      : hashlib.sha1,
            "sha256"    : hashlib.sha256,
        }

        self.Parsing_Function = {
            "filename"  : self.getFileName,
            "size"      : lambda file : f"{ format(os.stat(file)[6], ',') } Byte",  # size of file
            "md5"       : lambda *arg : self.getHash(*arg),
            "sha1"      : lambda *arg : self.getHash(*arg),
            "sha256"    : lambda *arg : self.getHash(*arg),
            "ctime"     : lambda file : datetime.strptime(time.ctime(os.stat(file)[9]), "%a %b %d %H:%M:%S %Y").isoformat(),
            "mtime"     : lambda file : datetime.strptime(time.ctime(os.stat(file)[8]), "%a %b %d %H:%M:%S %Y").isoformat(), 
            "atime"     : lambda file : datetime.strptime(time.ctime(os.stat(file)[7]), "%a %b %d %H:%M:%S %Y").isoformat(), 
        }
        ##############################################


        self.doc = docx.Document()
        self.options = options
        

    def scanDir(self, dir, option = 'not_recursive') :
        '''
            Parsing files after scanning directories.
            
            @dir = Directory path. ex) 'sample/dir'
            @option value is only 'recursive' or 'not_recursive' 
        '''
        if option == 'not_recursive' :
            for _path in os.listdir(dir) :
                path = f'{dir}{self.path_delimiter}{_path}'
                if os.path.isfile(path):
                    self.insertFile(path)

        elif option == 'recursive' :
            for dir, _, files in os.walk(dir) :
                for file in files :
                    path = f'{dir}{self.path_delimiter}{file}'
                    if os.path.isfile(path) :
                        self.insertFile(path)

        else :
            return self.ErrorHandler(self.NOT_SUPPORTED_OPTION)

    
    def ErrorHandler(self, Message) :
        print(Message)
        return False
    
    def insertFile(self, path) :
        '''
            Write a new table to docx, add a file information
            
            @dir = Directory path. ex) 'sample/dir'
            @option value is only 'recursive' or 'not_recursive' 
        '''
        table = self.doc.add_table(
            rows = len(self.options), 
            cols = 2
        )

        table.style = 'Table Grid'
        
        for row in range(len(self.options)) :
            table.rows[row].cells[0].text = self.options[row]
            table.rows[row].cells[1].text = self.dataHandler(self.options[row], path)
        
        self.doc.add_paragraph()
    
    def getFileName(self, path) :
        return path.split(self.path_delimiter)[-1]

    
    def dataHandler(self, data, path) :
        if data in self.Hash_Algorithm.keys() :
            return self.Parsing_Function[data](path, data)
        return self.Parsing_Function[data](path)
        
    
    def getHash(self, path, Type) :
        hash = self.Hash_Algorithm[Type]()
        with open(path, 'rb') as f:
            hash.update(f.read())
        return hash.hexdigest()
    
    def save(self, path) :
        self.doc.save(path)